"""Compile approved chapters into a DOCX manuscript."""
from __future__ import annotations
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from . import config

from . import db

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False


def _safe_filename(s: str) -> str:
    s = re.sub(r"[^A-Za-z0-9_\- ]+", "", s).strip().replace(" ", "_")
    return s[:80] or "book"


def compile_book(book_id: str) -> str:
    """Compile approved chapters into a .docx; return path."""
    book = db.get_book(book_id)
    if not book:
        raise RuntimeError(f"Book {book_id} not found")
    chapters = db.list_chapters(book_id)
    if not chapters:
        raise RuntimeError("No chapters to compile")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)

    # Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(book["title"])
    run.bold = True
    run.font.size = Pt(28)
    doc.add_paragraph()
    doc.add_paragraph()
    if book.get("pre_notes"):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(book["pre_notes"][:200])
        sr.italic = True
    doc.add_page_break()

    # Chapters
    for ch in chapters:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = h.add_run(f"Chapter {ch['chapter_number']}")
        hr.bold = True
        hr.font.size = Pt(16)

        t2 = doc.add_paragraph()
        t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = t2.add_run(ch.get("title") or "")
        tr.italic = True
        tr.font.size = Pt(14)
        doc.add_paragraph()

        for para in (ch.get("content") or "").split("\n\n"):
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph(para)
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    os.makedirs(config.OUTPUT_DIR, exist_ok=True)
    fname = f"{_safe_filename(book['title'])}_{book_id[:8]}.docx"
    path = os.path.join(config.OUTPUT_DIR, fname)
    doc.save(path)
    
    # Also generate PDF if reportlab is available
    if HAS_REPORTLAB:
        try:
            compile_book_to_pdf(book_id, path)
        except Exception as e:
            print(f"[compiler] PDF generation failed: {e}")
    
    return path


def compile_book_to_pdf(book_id: str, docx_path: str = None) -> str:
    """Compile approved chapters into a PDF using reportlab."""
    if not HAS_REPORTLAB:
        raise ImportError("reportlab not installed. Install with: pip install reportlab")
    
    book = db.get_book(book_id)
    if not book:
        raise RuntimeError(f"Book {book_id} not found")
    chapters = db.list_chapters(book_id)
    if not chapters:
        raise RuntimeError("No chapters to compile")

    os.makedirs(config.OUTPUT_DIR, exist_ok=True)
    fname = f"{_safe_filename(book['title'])}_{book_id[:8]}.pdf"
    pdf_path = os.path.join(config.OUTPUT_DIR, fname)

    # Create PDF
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=28,
        textColor='black',
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=14,
        textColor='black',
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    
    chapter_title_style = ParagraphStyle(
        'ChapterTitle',
        parent=styles['Heading2'],
        fontSize=16,
        textColor='black',
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    chapter_subtitle_style = ParagraphStyle(
        'ChapterSubtitle',
        parent=styles['Normal'],
        fontSize=12,
        textColor='black',
        spaceAfter=15,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        textColor='black',
        spaceAfter=12,
        alignment=TA_LEFT,
        leading=16
    )

    story = []

    # Title page
    story.append(Paragraph(book["title"], title_style))
    story.append(Spacer(1, 0.3 * inch))
    if book.get("pre_notes"):
        story.append(Paragraph(book["pre_notes"][:200], subtitle_style))
    story.append(PageBreak())

    # Chapters
    for ch in chapters:
        story.append(Paragraph(f"Chapter {ch['chapter_number']}", chapter_title_style))
        story.append(Paragraph(ch.get("title") or "", chapter_subtitle_style))
        story.append(Spacer(1, 0.2 * inch))
        
        content = ch.get("content") or ""
        for para in content.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            story.append(Paragraph(para, body_style))
        
        story.append(PageBreak())

    doc.build(story)
    return pdf_path
